import io
from app.utils.sanitizer import sanitize_text

def extract_text_from_pdf(file_bytes: bytes) -> str:
    """
    Extracts text from PDF file bytes using pdfplumber with pypdf fallback.
    """
    text = ""
    try:
        import pdfplumber
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            pages_text = []
            for page in pdf.pages:
                extracted = page.extract_text()
                if extracted:
                    pages_text.append(extracted)
            text = "\n\n".join(pages_text)
    except Exception as e:
        print(f"pdfplumber extraction warning: {e}, falling back to pypdf...")
        try:
            from pypdf import PdfReader
            reader = PdfReader(io.BytesIO(file_bytes))
            pages_text = []
            for page in reader.pages:
                extracted = page.extract_text()
                if extracted:
                    pages_text.append(extracted)
            text = "\n\n".join(pages_text)
        except Exception as e2:
            raise ValueError(f"Failed to extract text from PDF: {e2}")

    if not text.strip():
        raise ValueError("Could not extract readable text from PDF. The file may be empty, image-only/scanned, or encrypted.")
        
    return sanitize_text(text)

def extract_text_from_docx(file_bytes: bytes) -> str:
    """
    Extracts text from DOCX file bytes using python-docx.
    """
    try:
        import docx
        doc = docx.Document(io.BytesIO(file_bytes))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        
        # Also extract table text
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join([cell.text.strip() for cell in row.cells if cell.text.strip()])
                if row_text:
                    paragraphs.append(row_text)
                    
        text = "\n".join(paragraphs)
    except Exception as e:
        raise ValueError(f"Failed to extract text from Word document (.docx): {e}")
        
    if not text.strip():
        raise ValueError("The uploaded Word document contains no readable text.")
        
    return sanitize_text(text)

def parse_resume_document(file_bytes: bytes, filename: str) -> str:
    """
    Routes document byte stream to appropriate parser based on extension.
    """
    ext = "." + filename.split(".")[-1].lower() if "." in filename else ""
    if ext == ".pdf":
        return extract_text_from_pdf(file_bytes)
    elif ext == ".docx":
        return extract_text_from_docx(file_bytes)
    else:
        raise ValueError(f"Unsupported file extension: {ext}")
